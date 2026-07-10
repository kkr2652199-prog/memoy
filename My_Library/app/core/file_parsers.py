"""파일 형식별 텍스트 추출 모듈.

분리된 서브모듈에서 re-export하여 기존 import 경로를 유지한다.
"""

import csv
import io
import re
from pathlib import Path

# ──────── 서브모듈 re-export (기존 import 호환성 유지) ────────
from app.core.url_detectors import (                     # noqa: F401
    GOOGLE_DOC_PATTERNS,
    YOUTUBE_PATTERNS,
    detect_google_doc_url,
    detect_youtube_url,
    detect_url_type,
)
from app.core.brand_extractor import (                   # noqa: F401
    DOMAIN_BRAND_MAP,
    BLOG_PLATFORMS,
    SNS_PLATFORMS,
    _DOMAIN_LABELS_FALLBACK,
    _normalize_host,
    _host_matches_map_key,
    _lookup_domain_brand,
    _lookup_blog_platform,
    _lookup_sns_platform,
    _lookup_domain_labels_fallback,
    _fetch_html_for_brand,
    _extract_og_site_name,
    _extract_naver_blog_nickname,
    _clean_naver_blog_nick,
    platform_hint_from_source,
    platform_code_to_korean,
    extract_source_brand,
)
from app.core.url_fetchers import (                      # noqa: F401
    build_google_export_url,
    fetch_google_doc,
    fetch_google_sheets_xlsx,
    _bs4_extract_body,
    fetch_webpage,
    fetch_youtube_transcript,
)

# ──────── 상수 ────────

SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".pdf", ".xlsx", ".xls", ".csv",
    ".docx", ".jpg", ".jpeg", ".png", ".html", ".htm",
}


# ──────── 엑셀/CSV 파싱 ────────

def _parse_excel_bytes(data: bytes) -> str:
    """xlsx 바이트에서 모든 시트의 텍스트를 추출한다."""
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
        tmp.write(data)
        tmp_path = Path(tmp.name)
    try:
        return _parse_excel(tmp_path)
    finally:
        tmp_path.unlink(missing_ok=True)


def get_excel_sheets_info(filepath_or_bytes) -> list[dict]:
    """엑셀 파일에서 시트별 이름과 행 수를 반환한다."""
    from openpyxl import load_workbook
    import tempfile

    if isinstance(filepath_or_bytes, bytes):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
            tmp.write(filepath_or_bytes)
            fp = Path(tmp.name)
        cleanup = True
    else:
        fp = Path(filepath_or_bytes)
        cleanup = False

    try:
        wb = load_workbook(str(fp), read_only=True, data_only=True)
        sheets = []
        for name in wb.sheetnames:
            ws = wb[name]
            row_count = 0
            for _ in ws.iter_rows(values_only=True):
                row_count += 1
            sheets.append({"name": name, "row_count": row_count})
        wb.close()
        return sheets
    finally:
        if cleanup:
            fp.unlink(missing_ok=True)


def extract_single_sheet_text(filepath_or_bytes, sheet_name: str) -> str:
    """엑셀 파일에서 특정 시트의 텍스트만 추출한다."""
    from openpyxl import load_workbook
    import tempfile

    if isinstance(filepath_or_bytes, bytes):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
            tmp.write(filepath_or_bytes)
            fp = Path(tmp.name)
        cleanup = True
    else:
        fp = Path(filepath_or_bytes)
        cleanup = False

    try:
        wb = load_workbook(str(fp), read_only=True, data_only=True)
        if sheet_name not in wb.sheetnames:
            wb.close()
            return f"(시트 '{sheet_name}'를 찾을 수 없습니다)"
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        wb.close()
        return f"[시트: {sheet_name}]\n" + "\n".join(rows) if rows else f"(시트 '{sheet_name}'에 데이터가 없습니다)"
    finally:
        if cleanup:
            fp.unlink(missing_ok=True)


# ──────── 파일 추출 ────────

def extract_text_from_file(filepath: Path) -> str:
    ext = filepath.suffix.lower()
    if ext in (".txt", ".md"):
        return _parse_text(filepath)
    elif ext == ".pdf":
        return _parse_pdf(filepath)
    elif ext in (".xlsx", ".xls"):
        return _parse_excel(filepath)
    elif ext == ".csv":
        return _parse_csv(filepath)
    elif ext == ".docx":
        return _parse_docx(filepath)
    elif ext in (".jpg", ".jpeg", ".png"):
        return _parse_image(filepath)
    elif ext in (".html", ".htm"):
        return _parse_html(filepath)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {ext}")


def get_file_type_label(ext: str) -> str:
    labels = {
        ".txt": "텍스트", ".md": "마크다운", ".pdf": "PDF",
        ".xlsx": "엑셀", ".xls": "엑셀", ".csv": "CSV",
        ".docx": "워드", ".jpg": "이미지(OCR)", ".jpeg": "이미지(OCR)",
        ".png": "이미지(OCR)", ".html": "HTML", ".htm": "HTML",
    }
    return labels.get(ext.lower(), "기타")


def _parse_text(filepath: Path) -> str:
    for encoding in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            return filepath.read_text(encoding=encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return filepath.read_text(encoding="utf-8", errors="replace")


def _parse_pdf(filepath: Path) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(str(filepath))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[페이지 {i + 1}]\n{text.strip()}")
    return "\n\n".join(pages) if pages else "(PDF에서 텍스트를 추출할 수 없습니다)"


def _parse_excel(filepath: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(filepath), read_only=True, data_only=True)
    sheets_text = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            header = f"[시트: {sheet_name}]"
            sheets_text.append(f"{header}\n{chr(10).join(rows)}")
    wb.close()
    return "\n\n".join(sheets_text) if sheets_text else "(엑셀에서 데이터를 추출할 수 없습니다)"


def get_tabular_preview(filepath: Path, max_rows_per_sheet: int = 80) -> list[dict] | None:
    """엑셀/CSV 원본을 시트별 행 배열로 반환한다 (UI 표 미리보기용)."""
    ext = filepath.suffix.lower()
    if ext == ".csv":
        return _tabular_preview_csv(filepath, max_rows_per_sheet)
    if ext in (".xlsx", ".xlsm"):
        return _tabular_preview_xlsx(filepath, max_rows_per_sheet)
    if ext == ".xls":
        return None
    return None


def _tabular_preview_csv(filepath: Path, max_rows: int) -> list[dict]:
    encodings = ("utf-8", "cp949", "euc-kr", "latin-1")
    rows_out: list[list[str]] = []
    for enc in encodings:
        try:
            with open(filepath, encoding=enc, errors="replace", newline="") as f:
                reader = csv.reader(f)
                for i, row in enumerate(reader):
                    if i >= max_rows:
                        break
                    rows_out.append([str(c) if c is not None else "" for c in row])
            break
        except UnicodeDecodeError:
            rows_out = []
            continue
    if not rows_out:
        return []
    return [{"name": "Sheet1", "rows": rows_out}]


def _tabular_preview_xlsx(filepath: Path, max_rows_per_sheet: int) -> list[dict]:
    from openpyxl import load_workbook

    wb = load_workbook(str(filepath), read_only=True, data_only=True)
    result = []
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= max_rows_per_sheet:
                    break
                rows.append(["" if c is None else str(c) for c in row])
            result.append({"name": sheet_name, "rows": rows})
    finally:
        wb.close()
    return result


def _parse_csv(filepath: Path) -> str:
    text = _parse_text(filepath)
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        if any(cell.strip() for cell in row):
            rows.append(" | ".join(row))
    return "\n".join(rows) if rows else "(CSV에서 데이터를 추출할 수 없습니다)"


def _parse_docx(filepath: Path) -> str:
    from docx import Document

    doc = Document(str(filepath))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            paragraphs.append("\n".join(rows))

    return "\n\n".join(paragraphs) if paragraphs else "(워드 문서에서 텍스트를 추출할 수 없습니다)"


def _parse_image(filepath: Path) -> str:
    from PIL import Image

    try:
        import pytesseract
    except ImportError:
        return "(pytesseract가 설치되지 않았습니다. OCR을 사용하려면 Tesseract를 설치하세요.)"

    try:
        img = Image.open(str(filepath))
        text = pytesseract.image_to_string(img, lang="kor+eng")
        return text.strip() if text.strip() else "(이미지에서 텍스트를 추출할 수 없습니다)"
    except Exception as e:
        return f"(OCR 실패: {e}. Tesseract가 설치되어 있는지 확인하세요.)"


def _parse_html(filepath: Path) -> str:
    from bs4 import BeautifulSoup

    raw = _parse_text(filepath)
    soup = BeautifulSoup(raw, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines) if lines else "(HTML에서 텍스트를 추출할 수 없습니다)"
